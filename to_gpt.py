import os
from openai import OpenAI
from docx import Document
import json

# Initialize the OpenAI client
client = OpenAI(api_key='')

# Function to extract text from a .docx file
def extract_text_from_docx(file_path):
    """
    Извлекает текст из документа Word, включая текст из параграфов, таблиц и других элементов.
    """
    doc = Document(file_path)
    full_text = []

    # Извлечение текста из параграфов
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Пропускаем пустые строки
            full_text.append(paragraph.text)

    # Извлечение текста из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():  # Пропускаем пустые строки
                        full_text.append(paragraph.text)

    # Возвращаем текст, соединённый через перенос строки
    return '\n'.join(full_text)

# Function to send content to GPT-4 for JSON generation
def send_to_gpt4_for_json(content, model="gpt-4o-mini-2024-07-18", max_tokens=3000):
    try:
        prompt_text = (
            "Преобразуй следующий текст теста в JSON, не изменяя его содержимое (сохраняй все пробелы, знаки препинания и форматирование).\n\n"
            "Используй следующую структуру:\n"
            "```\n"
            "{\n"
            "    \"title\": \"Название теста\",\n"
            "    \"questions\": [\n"
            "        {\n"
            "            \"number\": номер вопроса,\n"
            "            \"question\": \"текст вопроса\",\n"
            "            \"options\": [\"вариант 1\", \"вариант 2\", ...],\n"
            "            \"answer\": \"правильный ответ(ы)\"\n"
            "        },\n"
            "        ...\n"
            "    ]\n"
            "}\n"
            "```\n\n"
            "Инструкции:\n"
            "1. Извлеки из текста название теста и помести его в поле \"title\".\n"
            "2. Каждый вопрос должен содержать порядковый номер, полный текст вопроса, массив вариантов ответа (если они присутствуют) и поле для правильного ответа.\n"
            "3. Если в вопросе присутствует раздел с ответами, начинающийся с \"Ответы:\" или \"Запишите ответ:\", извлеки его целиком и помести в поле \"answer\".\n"
            "4. Если вариантов ответа нет, оставь поле \"options\" пустым массивом [].\n"
            "5. Если правильный ответ не указан, оставь поле \"answer\" пустым.\n\n"
            "----------------------------------------\n"
            "Вот текст для преобразования:"
        )
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Ты помощник, который преобразует текст тестов в JSON строгой структуры. "
                        "Не изменяй исходное содержимое текста."
                    )
                },
                {
                    "role": "user",
                    "content": prompt_text
                },
                {
                    "role": "user",
                    "content": content
                }
            ],
            max_tokens=max_tokens,
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error contacting GPT-4 API: {e}")
        return ""


# Function to validate and enforce the JSON structure
def validate_and_fix_json(json_data):
    try:
        # Check if the main keys exist
        if not isinstance(json_data, dict) or "title" not in json_data or "questions" not in json_data:
            raise ValueError("Invalid JSON structure")
        
        # Validate questions
        if not isinstance(json_data["questions"], list):
            json_data["questions"] = []
        
        for question in json_data["questions"]:
            if not isinstance(question, dict):
                continue
            # Ensure required keys exist
            question.setdefault("number", 0)
            question.setdefault("question", "")
            question.setdefault("options", [])
            question.setdefault("answer", "")
        
        return json_data
    except Exception as e:
        print(f"Validation error: {e}")
        # Return an empty JSON with the correct structure if validation fails
        return {
            "title": "",
            "questions": []
        }

# Function to save parsed data to a .json file
def save_parsed_data_to_json(parsed_data, output_file):
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(parsed_data, f, ensure_ascii=False, indent=4)
        print(f"Parsed data saved to {output_file}")
    except Exception as e:
        print(f"Error saving parsed data: {e}")

# Function to process a single file
def process_file(file_path):
    try:
    
        json_file_path = file_path.replace(".docx", ".json")

        # Check if the JSON file already exists
        if os.path.exists(json_file_path):
            print(f"JSON file already exists for {file_path}, skipping.")
            return  # Skip processing this file
            
        print(f"Processing file: {file_path}")
        
        # Extract text
        content = extract_text_from_docx(file_path)
        
        # Ask GPT-4 to generate JSON
        gpt_response = send_to_gpt4_for_json(content)
        
        # Log raw response for debugging
        print(f"GPT-4 Response for {file_path}:\n{gpt_response}")
        
        # Parse GPT-4 response into JSON
        try:
            # Ensure the response is stripped of extra spaces and line breaks
            stripped_response = gpt_response.strip()
            
            # Remove any accidental code block delimiters (```)
            if stripped_response.startswith("```"):
                stripped_response = stripped_response.split("```json")[-1].split("```")[0]
            
            parsed_data = json.loads(stripped_response)
        except json.JSONDecodeError as e:
            print(f"Error decoding GPT response for {file_path}: {e}")
            parsed_data = {
                "title": "",
                "questions": []
            }
        
        # Validate and fix JSON structure
        validated_data = validate_and_fix_json(parsed_data)
        
        # Save to JSON file in the same directory
        json_file_path = file_path.replace(".docx", ".json")
        save_parsed_data_to_json(validated_data, json_file_path)
    
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")

# Main processing logic
if __name__ == "__main__":
    directory_to_process = "/mnt/ks/Works/3nd_tests/ready(last)"
    print(f"Walking through directory: {directory_to_process}")
    for root, _, files in os.walk(directory_to_process):
        for file in files:
            print(f"Found file: {file}")
            # Option 1: Process all DOCX files
            if file.endswith(".docx") and not file.startswith(".~lock"):
            # Option 2: Process DOCX files while excluding temporary lock files
            # if file.endswith(".docx") and not file.startswith(".~lock"):
                file_path = os.path.join(root, file)
                print(f"Processing file: {file_path}")
                process_file(file_path)
    print("Processing complete.")

#
