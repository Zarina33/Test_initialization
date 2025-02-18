#!/usr/bin/env python3
import os
import logging
from pathlib import Path
from docx import Document

# Настройка логирования: вывод в консоль и запись в файл
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("docx_extraction.log", encoding="utf-8"),
        logging.StreamHandler()
    ]
)

def extract_text_from_docx(file_path):
    """
    Извлекает текст из документа Word, включая текст из параграфов и таблиц.
    
    Args:
        file_path (str): Путь к файлу DOCX.
        
    Returns:
        str: Извлечённый текст.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logging.error(f"Ошибка открытия файла {file_path}: {e}")
        return ""
        
    full_text = []
    
    # Извлечение текста из параграфов
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Пропускаем пустые строки
            full_text.append(paragraph.text)
    
    # Извлечение текста из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
    
    return "\n".join(full_text)

def extract_text_from_directory(source_dir, output_dir):
    """
    Рекурсивно извлекает текст из всех файлов DOCX в директории source_dir
    и сохраняет результаты в output_dir, создавая аналогичную структуру директорий.
    
    Для каждого файла DOCX создаётся текстовый файл с расширением .txt,
    содержащий извлечённый текст.
    
    Args:
        source_dir (str): Путь к исходной директории с файлами DOCX.
        output_dir (str): Путь к директории, где будут сохранены текстовые файлы.
    """
    source_dir = Path(source_dir)
    output_dir = Path(output_dir)
    
    logging.info(f"Начало извлечения текста из DOCX файлов в {source_dir}")
    docx_files = list(source_dir.rglob("*.docx"))
    logging.info(f"Найдено {len(docx_files)} файлов DOCX")
    
    for docx_file in docx_files:
        try:
            text = extract_text_from_docx(docx_file)
            
            # Вычисляем относительный путь относительно исходной директории
            relative_path = docx_file.relative_to(source_dir)
            # Формируем путь для текстового файла (заменяем расширение .docx на .txt)
            output_file = output_dir / relative_path.with_suffix(".txt")
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)
            
            logging.info(f"Извлечён текст из {docx_file} -> {output_file}")
        except Exception as e:
            logging.error(f"Ошибка при обработке файла {docx_file}: {e}")
    
    logging.info("Извлечение текста завершено.")

if __name__ == "__main__":
    # Укажите исходную директорию с DOCX файлами
    source_directory = "/mnt/ks/Works/3nd_tests/errors_folder"
    # Укажите директорию для сохранения извлечённого текста
    output_directory = "/mnt/ks/Works/3nd_tests/extracted_text"
    
    extract_text_from_directory(source_directory, output_directory)
    print("Извлечение текста завершено.")

