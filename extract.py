import os
import re
import shutil
import xml.etree.ElementTree as ET
from docx import Document
from lxml import etree
from urllib.parse import unquote

def check_docx_content(source_path):
    try:
        doc = Document(source_path)
        doc_xml = doc._element.xml
        has_images = len(doc.inline_shapes) > 0
        has_math = '<m:oMath' in doc_xml
        return has_math, has_images
    except Exception as e:
        print(f"Error checking document content: {str(e)}")
        return False, False

def replace_content_with_paths(source_path, dest_path):
    try:
        doc = Document(source_path)
        doc_name = os.path.splitext(os.path.basename(source_path))[0]
        base_dir = os.path.dirname(dest_path)
        extracted_base_dir = os.path.join(base_dir, f"extracted_files_{doc_name}")

        math_dir = os.path.join(extracted_base_dir, "math_files")
        images_dir = os.path.join(extracted_base_dir, "images")

        os.makedirs(math_dir, exist_ok=True)
        os.makedirs(images_dir, exist_ok=True)

        modified = False

        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
        }

        # Process images
        for i, shape in enumerate(doc.inline_shapes, 1):
            try:
                if shape.type == 3:  # Picture
                    image_filename = f"{doc_name}_image_{i}.png"
                    image_path = os.path.join(images_dir, image_filename)

                    image_part = shape._inline.graphic.graphicData.pic.blipFill.blip
                    image_rId = image_part.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed']
                    image_part = doc.part.related_parts[image_rId]

                    with open(image_path, 'wb') as f:
                        f.write(image_part.blob)

                    for paragraph in doc.paragraphs:
                        if shape._inline in paragraph._element.iter():
                            paragraph.text = f"[Изображение заменено: {image_path}]"

                    print(f"Replaced image {i} with path: {image_path}")
                    modified = True
            except Exception as e:
                print(f"Error processing image {i}: {str(e)}")

        # Process math formulas
        math_formulas = doc._element.xpath('.//*[local-name()="oMath" or local-name()="oMathPara"]')
        for i, math_formula in enumerate(math_formulas, 1):
            try:
                math_filename = f"{doc_name}_math_{i}.xml"
                math_path = os.path.join(math_dir, math_filename)

                with open(math_path, 'w', encoding='utf-8') as f:
                    f.write(etree.tostring(math_formula, encoding='unicode', pretty_print=True))

                parent = math_formula.getparent()
                while parent is not None and etree.QName(parent).localname != 'p':
                    parent = parent.getparent()

                if parent is not None:
                    try:
                        parent.remove(math_formula)
                    except Exception as inner_e:
                        print(f"Warning: couldn't remove math formula node: {inner_e}")

                    new_r = etree.Element(f"{{{ns['w']}}}r")
                    new_t = etree.Element(f"{{{ns['w']}}}t")
                    new_t.text = f"[Формула заменена: {math_path}]"
                    new_r.append(new_t)
                    parent.append(new_r)

                print(f"Replaced math formula {i} with path: {math_path}")
                modified = True
            except Exception as e:
                print(f"Error processing math formula {i}: {str(e)}")

        if modified:
            doc.save(dest_path)
            print(f"Saved modified document with path references to: {dest_path}")
        else:
            shutil.copy2(source_path, dest_path)
            print(f"No modifications needed, copied original file to: {dest_path}")

    except Exception as e:
        raise Exception(f"Error processing document: {str(e)}")

def process_docx(source_path, destination_path):
    has_math, has_images = check_docx_content(source_path)

    if not (has_math or has_images):
        print(f"No math formulas or images found in {source_path}")
        os.makedirs(os.path.dirname(destination_path), exist_ok=True)
        shutil.copy2(source_path, destination_path)
        print(f"Copied file to {destination_path}")
        return

    print(f"Found content to process in {source_path}:")
    if has_math: print("- Math formulas")
    if has_images: print("- Images")

    replace_content_with_paths(source_path, destination_path)

def process_directory(source_dir, destination_dir):
    for root, _, files in os.walk(source_dir):
        for file in files:
            if file.endswith('.docx'):
                source_path = os.path.join(root, file)
                rel_path = os.path.relpath(root, source_dir)
                dest_dir = os.path.join(destination_dir, rel_path)
                os.makedirs(dest_dir, exist_ok=True)
                destination_path = os.path.join(dest_dir, file)
                
                try:
                    print(f"\nProcessing: {source_path}")
                    process_docx(source_path, destination_path)
                except Exception as e:
                    print(f"Error processing {source_path}: {str(e)}")

source_dir = "/mnt/ks/Works/3nd_tests/ToBeResized/Геометрия 10 класс/Русская версия/S-10-003"
destination_dir = "/mnt/ks/Works/3nd_tests/new"

process_directory(source_dir, destination_dir)
print("\nBatch processing completed")