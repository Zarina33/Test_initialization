import os
from docx import Document
from shutil import copy2

def convert_tables_to_text(file_path):
    print(f"Начинаем обработку файла: {file_path}")
    doc = Document(file_path)
    new_doc = Document()
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            new_doc.add_paragraph(paragraph.text)
    
    table_count = len(doc.tables)
    print(f"Найдено таблиц в документе: {table_count}")
    
    for i, table in enumerate(doc.tables, 1):
        print(f"Обработка таблицы {i} из {table_count}")
        variants = {}
        
        for row in table.rows:
            cells = [cell.text.strip().lower() for cell in row.cells]
            if not cells or not any(cells):
                continue
            
            variant = cells[0].replace('вар.', '').strip()
            answers = cells[1:]
            
            if variant and answers:
                variants[variant] = answers
        
        new_doc.add_paragraph()
        
        for var_num, answers in sorted(variants.items()):
            var_paragraph = new_doc.add_paragraph(f"{var_num}-вар.")
            formatted_answers = [f"{i+1}){ans};" for i, ans in enumerate(answers) if ans]
            answers_text = " ".join(formatted_answers)
            new_doc.add_paragraph(answers_text)
            new_doc.add_paragraph()
    
    return new_doc

def process_file(file_path):
    try:
        print(f"\nОбработка файла: {file_path}")
        
        # Проверяем размер файла
        file_size = os.path.getsize(file_path)
        print(f"Размер файла: {file_size} байт")
        
        # Создаем имя для старого файла
        file_dir = os.path.dirname(file_path)
        file_name = os.path.basename(file_path)
        file_base = os.path.splitext(file_name)[0]
        old_file_path = os.path.join(file_dir, f"{file_base} (Old).docx")
        
        # Копируем оригинальный файл
        print(f"Создаем резервную копию: {os.path.basename(old_file_path)}")
        copy2(file_path, old_file_path)
        
        # Конвертируем таблицы в текст
        print("Конвертируем таблицы в текст...")
        new_doc = convert_tables_to_text(file_path)
        
        # Сохраняем новый документ
        print(f"Сохраняем изменения в: {file_name}")
        new_doc.save(file_path)
        
        print(f"Успешно обработан файл: {file_name}")
        return True
        
    except Exception as e:
        print(f"Ошибка при обработке файла {file_path}: {str(e)}")
        return False

def process_directory_recursive(root_dir):
    print(f"Начинаем обработку директории: {root_dir}")
    processed_files = 0
    failed_files = 0
    
    for current_dir, dirs, files in os.walk(root_dir):
        print(f"\nПросматриваем папку: {current_dir}")
        docx_files = [f for f in files if f.endswith('.docx')]
        
        if docx_files:
            print(f"Найдено DOCX файлов в текущей папке: {len(docx_files)}")
            for file_name in docx_files:
                file_path = os.path.join(current_dir, file_name)
                if process_file(file_path):
                    processed_files += 1
                else:
                    failed_files += 1
    
    return processed_files, failed_files

if __name__ == "__main__":
    folder_path = "/mnt/ks/Works/3nd_tests/tables/Геометрия 11 класс"
    
    if not folder_path:
        folder_path = os.getcwd()
        print(f"Используем текущую папку: {folder_path}")
    
    if os.path.exists(folder_path):
        processed, failed = process_directory_recursive(folder_path)
        print("\nОбработка завершена!")
        print(f"Успешно обработано файлов: {processed}")
        if failed > 0:
            print(f"Не удалось обработать файлов: {failed}")
    else:
        print(f"Указанная папка не существует: {folder_path}")




